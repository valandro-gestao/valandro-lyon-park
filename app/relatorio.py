"""Gerador de relatórios Word (.docx) com logo e tabelas."""
import os, io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models import ResultadoUnidade, get_historico_anual
from app.calculators.patio import ResultadoPatio
from app.engine import get_unit

LOGO_PATH = os.path.join(os.path.dirname(__file__), "../assets/logo.png")
OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "../output")

COR_VERDE = RGBColor(0x00, 0xB0, 0xA0)
COR_CINZA = RGBColor(0xF2, 0xF2, 0xF2)
COR_TEXTO = RGBColor(0x26, 0x26, 0x26)

MESES_PT = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
            "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]


def fmt_br(v) -> str:
    if v is None or v == "":
        return "—"
    if isinstance(v, str):
        return v
    s = f"{abs(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"-{s}" if v < 0 else s


def fmt_pct(v: float) -> str:
    return f"{v*100:.2f}%"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_cell(cell, bold=True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold


def _new_doc() -> Document:
    doc = Document()
    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    # Fonte padrão
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].font.color.rgb = COR_TEXTO
    return doc


def _add_logo(doc: Document):
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].clear()
    run = header.paragraphs[0].add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Cm(5))
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_footer(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    footer.paragraphs[0].clear()
    run = footer.paragraphs[0].add_run(
        "LYON PARK ESTACIONAMENTOS    |    gestao@lyonpark.com.br    |    www.lyonpark.com.br"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_titulo(doc: Document, texto: str):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COR_VERDE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)


def _add_subtexto(doc: Document, texto: str):
    p = doc.add_paragraph(texto)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(10)


def _add_tabela_mensal(doc: Document, nome_bloco: str, meses: list[str],
                        linhas_dados: list[tuple], is_header_gray=True):
    """
    linhas_dados: list of (label: str, valores: list[str | float])
    """
    n_cols = 1 + len(meses)
    table = doc.add_table(rows=1 + len(linhas_dados), cols=n_cols)
    table.style = "Table Grid"

    # Cabeçalho
    header_row = table.rows[0]
    _set_cell_bg(header_row.cells[0], "00B0A0")
    run = header_row.cells[0].paragraphs[0].add_run(nome_bloco)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)

    for i, mes in enumerate(meses):
        cell = header_row.cells[i + 1]
        _set_cell_bg(cell, "00B0A0")
        run = cell.paragraphs[0].add_run(mes)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Linhas de dados
    for row_idx, (label, valores) in enumerate(linhas_dados):
        row = table.rows[row_idx + 1]
        is_aluguel = "aluguel" in label.lower() or "repasse" in label.lower() or "total" in label.lower()
        bg = "E8F8F6" if is_aluguel else ("F9F9F9" if row_idx % 2 == 0 else "FFFFFF")

        _set_cell_bg(row.cells[0], bg)
        run = row.cells[0].paragraphs[0].add_run(label)
        run.font.size = Pt(9.5)
        if is_aluguel:
            run.bold = True

        for i, val in enumerate(valores):
            cell = row.cells[i + 1]
            _set_cell_bg(cell, bg)
            txt = fmt_br(val) if isinstance(val, (int, float)) else str(val)
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(9.5)
            if is_aluguel:
                run.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Largura das colunas
    table.columns[0].width = Cm(5.5)
    for i in range(1, n_cols):
        table.columns[i].width = Cm(2.8)

    doc.add_paragraph()
    return table


def _add_secao_historico(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("Histórico da Operação – Anos anteriores")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COR_VERDE
    p.paragraph_format.space_before = Pt(12)

    p2 = doc.add_paragraph("Indicadores consolidados para análise de desempenho e evolução da unidade.")
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p2.paragraph_format.space_after = Pt(6)


def gerar_relatorio_simples(resultado: ResultadoUnidade,
                              historico: list[dict] = None,
                              nome_arquivo: str = None) -> str:
    cfg = get_unit(resultado.unidade_id)
    linhas_cfg = cfg.get("relatorio", {}).get("linhas", [])
    ano = resultado.mes_referencia[:4]
    mes_idx = int(resultado.mes_referencia[5:7]) - 1

    doc = _new_doc()
    _add_logo(doc)
    _add_footer(doc)

    _add_titulo(doc, f"Prestação de contas de {ano} da unidade {cfg['contratante']}:")
    _add_subtexto(doc,
        "Abaixo resumo da operação. Para acesso a todas informações utilizar nossa "
        "plataforma de administração de estacionamento (e-Cloud) com login do contratante.")

    # Coletar meses já lançados no ano corrente (placeholder: só o mês atual)
    meses_disponiveis = MESES_PT[:mes_idx + 1]

    # Montar linhas da tabela — apenas as relevantes para esta unidade
    def linha(label, valores):
        return (label, valores)

    linhas_tabela = _montar_linhas_tabela(resultado, linhas_cfg, meses_disponiveis)
    _add_tabela_mensal(doc, cfg["nome"], meses_disponiveis, linhas_tabela)

    # Histórico anual
    if historico:
        _add_secao_historico(doc)
        anos = [str(h["ano"]) for h in historico]
        linhas_hist = _montar_linhas_historico(historico, linhas_cfg)
        _add_tabela_mensal(doc, cfg["nome"], anos, linhas_hist)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    nome = nome_arquivo or f"{cfg['id']}_{resultado.mes_referencia}.docx"
    caminho = os.path.join(OUTPUT_DIR, nome)
    doc.save(caminho)
    return caminho


def _montar_linhas_tabela(r: ResultadoUnidade, linhas_cfg: list, meses: list) -> list:
    """Gera linhas no formato [(label, [v1, v2, ...])] — por enquanto um único mês."""
    def v(val):
        return [val] + [""] * (len(meses) - 1)

    linhas = []
    mapa = {
        "faturamento":  ("Total Faturamento:", v(r.faturamento)),
        "aliquota":     ("Alíquota Imposto:", v(fmt_pct(r.aliquota_imposto))),
        "subtotal":     ("Subtotal:", v(r.subtotal)),
        "pe":           ("Ponto de Equilíbrio:", v(r.ponto_equilibrio)),
        "resultado":    ("Resultado:", v(r.resultado)),
        "prejuizo":     ("Prejuízo Acumulado", v(r.prejuizo_acumulado_entrada)),
        "aluguel":      ("Aluguel a Pagar", v(r.aluguel_calculado)),
        "taxa_admin":   ("Taxa de Administração", v(r.extras.get("taxa_admin", 0.0))),
        "adicional":    ("Adicional Fixo (48m)", v(r.extras.get("adicional_fixo", 0.0))),
    }
    for k, cv in r.custos.items():
        mapa[k] = (k.replace("_", " ").title(), v(cv))

    for l in linhas_cfg:
        if l in mapa:
            linhas.append(mapa[l])
    return linhas


def _montar_linhas_historico(historico: list, linhas_cfg: list) -> list:
    anos = [h for h in historico]

    def col(campo):
        return [h.get(campo, 0.0) for h in anos]

    mapa = {
        "faturamento":  ("Total Faturamento:", col("faturamento")),
        "subtotal":     ("Subtotal:", col("subtotal")),
        "pe":           ("Ponto de Equilíbrio:", col("ponto_equilibrio")),
        "resultado":    ("Resultado:", col("resultado")),
        "prejuizo":     ("Prejuízo Acumulado", col("prejuizo_acumulado_entrada")),
        "aluguel":      ("Aluguel Pago:", col("aluguel_calculado")),
    }
    linhas = []
    for l in linhas_cfg:
        if l in mapa:
            linhas.append(mapa[l])
    return linhas


def gerar_relatorio_patio(resultado: ResultadoPatio, split_id: str,
                           historico: list[dict] = None) -> str:
    """Gera relatório para Real ou Maiojama."""
    split_r = resultado.real if split_id == "real" else resultado.maiojama
    split_nome = "REAL (53,52%)" if split_id == "real" else "MAIOJAMA (46,48%)"
    mes = resultado.mes_referencia
    ano = mes[:4]
    mes_idx = int(mes[5:7]) - 1
    meses = MESES_PT[:mes_idx + 1]

    doc = _new_doc()
    _add_logo(doc)
    _add_footer(doc)

    _add_titulo(doc, f"Prestação de contas de {ano} da unidade Patio:")
    _add_subtexto(doc,
        "Abaixo resumo da operação. Para acesso a todas informações utilizar nossa "
        "plataforma de administração de estacionamento (e-Cloud) com login do contratante.")

    # Bloco operação estacionamento
    linhas_op = [
        ("Operação Estacionamento:", [split_r.faturamento] + [""] * (len(meses)-1)),
        ("Alíquota Imposto:", [fmt_pct(split_r.aliquota_imposto)] + [""] * (len(meses)-1)),
        ("Subtotal:", [split_r.subtotal] + [""] * (len(meses)-1)),
        ("Ponto de Equilíbrio:", [split_r.ponto_equilibrio] + [""] * (len(meses)-1)),
    ]
    for k, v in split_r.custos.items():
        linhas_op.append((k.replace("_", " ").title(), [v] + [""] * (len(meses)-1)))
    linhas_op += [
        ("Resultado:", [split_r.resultado] + [""] * (len(meses)-1)),
        ("Aluguel 95%:", [split_r.aluguel_calculado] + [""] * (len(meses)-1)),
    ]
    repasse_outros = split_r.extras.get("repasse_outros", 0.0)
    if repasse_outros:
        linhas_op.append(("Repasse Outros 50%:", [repasse_outros] + [""] * (len(meses)-1)))
        total = split_r.aluguel_calculado + repasse_outros
        linhas_op.append(("Total Repasse", [total] + [""] * (len(meses)-1)))

    _add_tabela_mensal(doc, split_nome, meses, linhas_op)

    # Histórico
    if historico:
        _add_secao_historico(doc)

    # Outros Serviços
    os_data = resultado.outros_servicos
    if os_data:
        p = doc.add_paragraph()
        p.add_run("Relatório de resultado de operações de Outros Serviços da unidade Trend Patio 24:").bold = True

        repasse_key = "repasse_real" if split_id == "real" else "repasse_maiojama"
        split_key = "REAL (53,52%)" if split_id == "real" else "MAIOJAMA (46,48%)"

        linhas_os = [
            ("Mídias", [os_data["receitas_midia"]] + [""] * (len(meses)-1)),
            ("Receitas Outros Serviços", [os_data["receitas_midia"]] + [""] * (len(meses)-1)),
            ("Alíquota Imposto:", [fmt_pct(0.1425)] + [""] * (len(meses)-1)),
            ("Subtotal:", [os_data["subtotal"]] + [""] * (len(meses)-1)),
        ]
        for k, v in os_data.get("despesas", {}).items():
            linhas_os.append((k.replace("_", " ").title(), [v] + [""] * (len(meses)-1)))
        linhas_os += [
            ("Total Despesas", [sum(os_data.get("despesas", {}).values())] + [""] * (len(meses)-1)),
            ("Resultado:", [os_data["resultado"]] + [""] * (len(meses)-1)),
            ("Saldo para Repasse", [os_data["resultado"]] + [""] * (len(meses)-1)),
            ("Repasse 50%", [os_data["repasse_total"]] + [""] * (len(meses)-1)),
            (split_key, [os_data[repasse_key]] + [""] * (len(meses)-1)),
        ]
        _add_tabela_mensal(doc, "TREND PATIO 24", meses, linhas_os)

    # Carregadores
    car = resultado.carregadores
    if car:
        p = doc.add_paragraph()
        p.add_run("Relatório de resultado de operação dos Carregadores Elétricos da unidade Trend Patio 24:").bold = True

        repasse_key = "repasse_real" if split_id == "real" else "repasse_maiojama"
        split_key = "REAL (53,52%)" if split_id == "real" else "MAIOJAMA (46,48%)"

        linhas_car = [
            ("Receita Carregadores", [car["receita"]] + [""] * (len(meses)-1)),
            ("Receitas Outros Serviços", [car["receita"]] + [""] * (len(meses)-1)),
            ("Taxa de Intermediação WEG", [car["taxa_weg"]] + [""] * (len(meses)-1)),
            ("Custo de Consumo Energia", [car["custo_energia"]] + [""] * (len(meses)-1)),
            ("Total Despesas", [car["taxa_weg"] + car["custo_energia"]] + [""] * (len(meses)-1)),
            ("Resultado:", [car["resultado"]] + [""] * (len(meses)-1)),
            ("Investimento Inicial", [car.get("investimento", 0.0)] + [""] * (len(meses)-1)),
            ("Saldo para Repasse", [car["saldo"]] + [""] * (len(meses)-1)),
            ("Repasse 60%", [car["repasse_total"]] + [""] * (len(meses)-1)),
            (split_key, [car[repasse_key]] + [""] * (len(meses)-1)),
        ]
        _add_tabela_mensal(doc, "TREND PATIO 24", meses, linhas_car)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    nome = f"patio_{split_id}_{mes}.docx"
    caminho = os.path.join(OUTPUT_DIR, nome)
    doc.save(caminho)
    return caminho


def gerar_relatorio_patio_manutencao(resultado: ResultadoPatio) -> str:
    man = resultado.manutencao
    mes = resultado.mes_referencia
    ano = mes[:4]
    mes_idx = int(mes[5:7]) - 1
    meses = MESES_PT[:mes_idx + 1]

    doc = _new_doc()
    _add_logo(doc)
    _add_footer(doc)

    _add_titulo(doc, f"Relatório de Manutenções de {ano} da unidade Patio:")

    linhas = [
        ("Receita de Manutenção", [man["receita"]] + [""] * (len(meses)-1)),
        ("Retenção de ISS - 5%", [man["retencao_iss"]] + [""] * (len(meses)-1)),
        ("Total Líquido", [man["total_liquido"]] + [""] * (len(meses)-1)),
        ("Total Despesas", [man["total_despesas"]] + [""] * (len(meses)-1)),
    ]
    for k, v in man.get("despesas", {}).items():
        linhas.append((k.replace("_", " ").title(), [v] + [""] * (len(meses)-1)))
    linhas += [
        ("Resultado", [man["resultado"]] + [""] * (len(meses)-1)),
        ("Saldo Acumulado", [man["saldo_acumulado"]] + [""] * (len(meses)-1)),
    ]
    _add_tabela_mensal(doc, "TREND PATIO 24", meses, linhas)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    caminho = os.path.join(OUTPUT_DIR, f"patio_manutencao_{mes}.docx")
    doc.save(caminho)
    return caminho


def gerar_todos(resultados: dict, mes: str) -> list[str]:
    arquivos = []
    for uid, resultado in resultados.items():
        if isinstance(resultado, ResultadoPatio):
            arquivos.append(gerar_relatorio_patio(resultado, "real"))
            arquivos.append(gerar_relatorio_patio(resultado, "maiojama"))
            if resultado.manutencao:
                arquivos.append(gerar_relatorio_patio_manutencao(resultado))
        else:
            historico = get_historico_anual(uid)
            arquivos.append(gerar_relatorio_simples(resultado, historico or None))
    return arquivos
